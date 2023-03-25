from lab7.tables_ import (
    initial_perm,
    exp_d,
    per,
    sbox,
    final_perm,
    keyp,
    shift_table,
    key_comp,
)

from docx import Document
import time
import pandas as pd


# Hexadecimal to binary conversion
def hex2bin(s):
    mp = {
        "0": "0000",
        "1": "0001",
        "2": "0010",
        "3": "0011",
        "4": "0100",
        "5": "0101",
        "6": "0110",
        "7": "0111",
        "8": "1000",
        "9": "1001",
        "A": "1010",
        "B": "1011",
        "C": "1100",
        "D": "1101",
        "E": "1110",
        "F": "1111",
    }
    bin = ""
    for i in range(len(s)):
        bin = bin + mp[s[i]]
    return bin


# Binary to hexadecimal conversion
def bin2hex(s):
    mp = {
        "0000": "0",
        "0001": "1",
        "0010": "2",
        "0011": "3",
        "0100": "4",
        "0101": "5",
        "0110": "6",
        "0111": "7",
        "1000": "8",
        "1001": "9",
        "1010": "A",
        "1011": "B",
        "1100": "C",
        "1101": "D",
        "1110": "E",
        "1111": "F",
    }
    hex = ""
    for i in range(0, len(s), 4):
        ch = ""
        ch = ch + s[i]
        ch = ch + s[i + 1]
        ch = ch + s[i + 2]
        ch = ch + s[i + 3]
        hex = hex + mp[ch]

    return hex


# Binary to decimal conversion
def bin2dec(binary):
    binary1 = binary
    decimal, i, n = 0, 0, 0
    while binary != 0:
        dec = binary % 10
        decimal = decimal + dec * pow(2, i)
        binary = binary // 10
        i += 1
    return decimal


# Decimal to binary conversion
def dec2bin(num):
    res = bin(num).replace("0b", "")
    if len(res) % 4 != 0:
        div = len(res) / 4
        div = int(div)
        counter = (4 * (div + 1)) - len(res)
        for i in range(0, counter):
            res = "0" + res
    return res


# Permute function to rearrange the bits
def permute(k, arr, n):
    permutation = ""
    for i in range(0, n):
        permutation = permutation + k[arr[i] - 1]
    return permutation


# shifting the bits towards left by nth shifts
def shift_left(k, nth_shifts):
    s = ""
    for i in range(nth_shifts):
        for j in range(1, len(k)):
            s = s + k[j]
        s = s + k[0]
        k = s
        s = ""
    return k


# calculating xow of two strings of binary number a and b
def xor(a, b):
    ans = ""
    for i in range(len(a)):
        if a[i] == b[i]:
            ans = ans + "0"
        else:
            ans = ans + "1"
    return ans


def encrypt(pt, rkb, rk):
    pt = hex2bin(pt)

    # Initial Permutation
    pt = permute(pt, initial_perm, 64)
    # print("After initial permutation", bin2hex(pt))

    # Splitting
    left = pt[0:32]
    right = pt[32:64]
    for i in range(0, 16):
        # Expansion D-box: Expanding the 32 bits data into 48 bits
        right_expanded = permute(right, exp_d, 48)

        # XOR RoundKey[i] and right_expanded
        xor_x = xor(right_expanded, rkb[i])

        # S-boxex: substituting the value from s-box table by calculating row and column
        sbox_str = ""
        for j in range(0, 8):
            row = bin2dec(int(xor_x[j * 6] + xor_x[j * 6 + 5]))
            col = bin2dec(
                int(
                    xor_x[j * 6 + 1]
                    + xor_x[j * 6 + 2]
                    + xor_x[j * 6 + 3]
                    + xor_x[j * 6 + 4]
                )
            )
            val = sbox[j][row][col]
            sbox_str = sbox_str + dec2bin(val)

        # Straight D-box: After substituting rearranging the bits
        sbox_str = permute(sbox_str, per, 32)

        # XOR left and sbox_str
        result = xor(left, sbox_str)
        left = result

        # Swapper
        if i != 15:
            left, right = right, left
        # print("Round ", i + 1, " ", bin2hex(left), " ", bin2hex(right), " ", rk[i])

    # Combination
    combine = left + right

    # Final permutation: final rearranging of bits to get cipher text
    cipher_text = permute(combine, final_perm, 64)
    return cipher_text


def main():
    stats = {}
    document = Document("medium.docx")
    print(f"Original text:\n{document.paragraphs[0].text}")
    pt = "".join([hex(ord(char))[2:].upper() for char in document.paragraphs[0].text])

    print(f"HEX text:\n{pt}")
    key = "AABB09182736CCDD"

    # Key generation
    # --hex to binary
    key = hex2bin(key)

    # getting 56 bit key from 64 bit using the parity bits
    key = permute(key, keyp, 56)

    # Splitting
    left = key[0:28]  # rkb for RoundKeys in binary
    right = key[28:56]  # rk for RoundKeys in hexadecimal

    rkb = []
    rk = []
    for i in range(0, 16):
        # Shifting the bits by nth shifts by checking from shift table
        left = shift_left(left, shift_table[i])
        right = shift_left(right, shift_table[i])

        # Combination of left and right string
        combine_str = left + right

        # Compression of key from 56 to 48 bits
        round_key = permute(combine_str, key_comp, 48)

        rkb.append(round_key)
        rk.append(bin2hex(round_key))

    print("\nEncryption")
    encryption_start = time.perf_counter()
    cipher_text = bin2hex(encrypt(pt, rkb, rk))
    stats["Encryption time"] = time.perf_counter() - encryption_start
    print("Cipher Text : ", cipher_text)

    print("\nDecryption")
    rkb_rev = rkb[::-1]
    rk_rev = rk[::-1]
    decryption_start = time.perf_counter()
    text = bin2hex(encrypt(cipher_text, rkb_rev, rk_rev))
    stats["Decryption time"] = time.perf_counter() - decryption_start
    print(f"Hex text:\n{text}")

    print("".join([chr(int(text[i : i + 2], 16)) for i in range(0, len(text), 2)]))
    print(stats)


if __name__ == "__main__":
    main()
